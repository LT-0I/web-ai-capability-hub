import asyncio, base64, hashlib, json, os, re, shutil, time, zipfile
from pathlib import Path
from urllib.parse import urlparse

from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError

RUN = Path('ip-literature-patent-research/.runs/literature-review-rl-antiuav-2026-05-13').resolve()
SCRIPTS = RUN / 'scripts'
SHOTS = RUN / 'round3-screenshots'
RAW = RUN / 'round3-raw-downloads'
SHOTS.mkdir(parents=True, exist_ok=True)
RAW.mkdir(parents=True, exist_ok=True)
URL = 'https://chatgpt.com/c/6a04a213-5648-83e8-b9d0-6134aef56831'
FINAL = RUN / 'round3-chatgpt-dr-export-strong-form.docx'
CANON = RUN / '强化学习在反无人机系统中的应用-文献综述.docx'
MANUAL = Path.home() / 'Downloads' / '强化学习在反无人机系统中的应用.docx'
LOG_PATH = SCRIPTS / 'round3-cdp-export-log.json'
SUSPICIOUS = re.compile(r'(export|download|docx|attachment|conversation|research|file|blob|download)', re.I)
DOCX_TEXT = re.compile(r'(DOCX|docx|Word|word|文档|下载\s*DOCX|下载|导出|Export|Download)')


def sha256(p: Path):
    if not p.exists(): return None
    h=hashlib.sha256()
    with p.open('rb') as f:
        for chunk in iter(lambda:f.read(1024*1024), b''):
            h.update(chunk)
    return h.hexdigest()


def looks_docx(path: Path):
    rec={'path': str(path), 'exists': path.exists()}
    if not path.exists(): return rec
    rec['size'] = path.stat().st_size
    try:
        from docx import Document
        d=Document(str(path))
        rec['paragraphs']=len(d.paragraphs)
        rec['chars']=sum(len(p.text) for p in d.paragraphs)
        rec['valid_docx']=True
    except Exception as e:
        rec['valid_docx']=False; rec['docx_error']=repr(e)
    try:
        with zipfile.ZipFile(path) as z:
            rec['zip_entries']=len(z.namelist())
            rec['has_word_document']='word/document.xml' in z.namelist()
    except Exception as e:
        rec['zip_error']=repr(e)
    rec['sha256']=sha256(path)
    rec['manual_sha256']=sha256(MANUAL)
    rec['byte_identical_to_manual'] = bool(rec.get('sha256') and rec.get('sha256') == rec.get('manual_sha256'))
    rec['substantive'] = rec.get('valid_docx') and rec.get('paragraphs',0) >= 50 and rec.get('chars',0) >= 5000
    return rec

async def screenshot(page, name, full=False):
    try:
        p=SHOTS/name
        await page.screenshot(path=str(p), full_page=full, timeout=20000)
        return str(p.relative_to(RUN))
    except Exception as e:
        return {'error':repr(e)}

async def pick_page(browser):
    ctx = browser.contexts[0]
    target = None
    for pg in ctx.pages:
        if pg.url.split('?')[0] == URL:
            target = pg; break
    if target is None:
        for pg in ctx.pages:
            if '/c/' in pg.url and 'chatgpt.com' in pg.url:
                target = pg; break
    if target is None:
        target = await ctx.new_page()
    if target.url.split('?')[0] != URL:
        await target.goto(URL, wait_until='domcontentloaded', timeout=60000)
    await target.bring_to_front()
    await target.set_viewport_size({'width':1500,'height':1000})
    await target.wait_for_timeout(6000)
    return target

async def frame_meta(page):
    out=[]
    for i,fr in enumerate(page.frames):
        rec={'index':i,'url':fr.url}
        try: rec['name']=fr.name
        except Exception: pass
        try: rec['title']=await fr.title()
        except Exception: pass
        try: rec['text']=(await fr.locator('body').inner_text(timeout=1500))[:2000]
        except Exception as e: rec['text_error']=repr(e)[:200]
        try:
            rec['elements']=await fr.evaluate('''() => Array.from(document.querySelectorAll('button,a,[role="button"],input')).map((el,idx)=>{const r=el.getBoundingClientRect(); const clean=s=>(s||'').replace(/\s+/g,' ').trim(); return {idx, tag:el.tagName.toLowerCase(), text:clean(el.innerText||el.textContent||''), aria:el.getAttribute('aria-label')||'', title:el.getAttribute('title')||'', href:el.getAttribute('href')||'', visible:!!(r.width&&r.height), box:{x:r.x,y:r.y,w:r.width,h:r.height}}})''')
        except Exception as e: rec['elements_error']=repr(e)[:200]
        out.append(rec)
    return out

async def find_candidates(page):
    cands=[]
    for i,fr in enumerate(page.frames):
        try:
            handles = await fr.locator('button,a,[role="button"],input').element_handles()
        except Exception:
            continue
        for idx,h in enumerate(handles):
            try:
                meta = await h.evaluate('''(el) => { const clean=s=>(s||'').replace(/\s+/g,' ').trim(); const r=el.getBoundingClientRect(); let p=el, ctx=''; for(let i=0;p&&i<5;i++,p=p.parentElement) ctx += ' || ' + clean(p.innerText||p.textContent||''); return {tag:el.tagName.toLowerCase(), text:clean(el.innerText||el.textContent||''), aria:el.getAttribute('aria-label')||'', title:el.getAttribute('title')||'', href:el.getAttribute('href')||'', download:el.getAttribute('download')||'', role:el.getAttribute('role')||'', visible:!!(r.width&&r.height), box:{x:r.x,y:r.y,w:r.width,h:r.height}, context:ctx.slice(0,1500)} }''')
                box = await h.bounding_box()
            except Exception:
                continue
            blob=' '.join(str(meta.get(k,'')) for k in ['text','aria','title','href','download','role','context'])
            score=0; low=blob.lower()
            if meta.get('visible'): score+=1
            if meta.get('aria') == '导出': score+=20
            if '下载 docx' in low: score+=18
            if 'docx' in low: score+=16
            if 'word' in low: score+=10
            if '导出' in blob or 'export' in low: score+=8
            if '下载' in blob or 'download' in low: score+=6
            if 'pdf' in low and 'docx' not in low: score-=5
            if '分享' in blob or 'share' in low: score-=5
            if score >= 8 and box and box.get('width',0) > 0 and box.get('height',0) > 0:
                cands.append({'score':score,'frame_index':i,'frame_url':fr.url,'idx':idx,'meta':meta,'bbox':box,'handle':h})
    cands.sort(key=lambda c:c['score'], reverse=True)
    return cands

async def raw_click(page_cdp, bbox, log, label):
    cx = bbox['x'] + bbox['width']/2
    cy = bbox['y'] + bbox['height']/2
    rec={'label':label,'x':cx,'y':cy,'bbox':bbox}
    log.setdefault('raw_clicks',[]).append(rec)
    await page_cdp.send('Input.dispatchMouseEvent', {'type':'mouseMoved','x':cx,'y':cy})
    await page_cdp.send('Input.dispatchMouseEvent', {'type':'mousePressed','x':cx,'y':cy,'button':'left','clickCount':1,'buttons':1})
    await page_cdp.send('Input.dispatchMouseEvent', {'type':'mouseReleased','x':cx,'y':cy,'button':'left','clickCount':1,'buttons':0})
    return rec

async def wait_for_download(download_state, timeout=60):
    start=time.time()
    while time.time()-start < timeout:
        for guid, rec in list(download_state.items()):
            if rec.get('state') == 'completed':
                path = RAW / guid
                rec['expected_path']=str(path)
                if path.exists():
                    return rec
            if rec.get('state') == 'canceled':
                continue
        # fallback: any fresh docx/crdownload done
        for p in RAW.iterdir():
            if p.is_file() and not p.name.endswith('.crdownload') and p.stat().st_size > 1000:
                return {'guid':p.name,'suggestedFilename':p.name,'state':'completed','expected_path':str(p), 'fallback_file_scan':True}
        await asyncio.sleep(0.5)
    return None

async def copy_if_valid(src: Path, log, tag):
    if not src.exists():
        log.setdefault('copy_attempts',[]).append({'tag':tag,'src':str(src),'exists':False})
        return False
    # If Chrome used guid name, still a zip/docx body.
    dest_raw = RAW / (tag + '-' + (src.name if src.suffix else src.name + '.docx'))
    if src.resolve() != dest_raw.resolve():
        try: shutil.copy2(src, dest_raw)
        except Exception: dest_raw = src
    shutil.copy2(src, FINAL)
    rec=looks_docx(FINAL); rec['tag']=tag; rec['src']=str(src); rec['raw_copy']=str(dest_raw)
    log.setdefault('copy_attempts',[]).append(rec)
    if rec.get('substantive') and not rec.get('byte_identical_to_manual'):
        shutil.copy2(FINAL, CANON)
        rec['canonical_overwritten']=True
        return True
    return False

async def click_format_items(page, page_cdp, log):
    # After a picker/menu opens, locate visible Word/DOCX/download items and raw-click them.
    await page.wait_for_timeout(1000)
    cands=await find_candidates(page)
    log['post_click_candidates']=[{k:v for k,v in c.items() if k!='handle'} for c in cands[:20]]
    for c in cands[:8]:
        blob=' '.join(str(c['meta'].get(k,'')) for k in ['text','aria','title','href','context']).lower()
        if ('docx' in blob or 'word' in blob or '下载 docx' in blob) and c['bbox']:
            await raw_click(page_cdp, c['bbox'], log, 'format/docx:'+str(c['meta'])[:160])
            await page.wait_for_timeout(2000)
            return True
    return False

async def main():
    log={'started':time.strftime('%Y-%m-%dT%H:%M:%S'), 'url':URL, 'manual_sha256':sha256(MANUAL), 'tactics':[]}
    async with async_playwright() as p:
        browser = await p.chromium.connect_over_cdp('http://127.0.0.1:9223')
        page = await pick_page(browser)
        page_cdp = await browser.contexts[0].new_cdp_session(page)
        browser_cdp = await browser.new_browser_cdp_session()
        log['page']={'url':page.url,'title':await page.title()}
        log['initial_screenshot']=await screenshot(page,'00-initial.png', full=True)
        log['initial_frames']=await frame_meta(page)

        download_state={}
        def on_will_begin(params):
            rec=dict(params); rec['event']='downloadWillBegin'; rec['ts']=time.time(); download_state[params.get('guid','')]=rec
            log.setdefault('browser_download_events',[]).append(rec)
        def on_progress(params):
            guid=params.get('guid','')
            download_state.setdefault(guid,{}).update(params)
            download_state[guid]['ts']=time.time()
            log.setdefault('browser_download_events',[]).append(dict(params, event='downloadProgress', ts=time.time()))
        browser_cdp.on('Browser.downloadWillBegin', on_will_begin)
        browser_cdp.on('Browser.downloadProgress', on_progress)
        try:
            await browser_cdp.send('Browser.setDownloadBehavior', {'behavior':'allowAndName','downloadPath':str(RAW),'eventsEnabled':True})
            log['set_download_behavior']='browser allowAndName ok'
        except Exception as e:
            log['set_download_behavior_error']=repr(e)
            try:
                await page_cdp.send('Browser.setDownloadBehavior', {'behavior':'allowAndName','downloadPath':str(RAW),'eventsEnabled':True})
                log['set_download_behavior']='page allowAndName ok'
            except Exception as e2:
                log['set_download_behavior_page_error']=repr(e2)

        # Network subscription for tactic 2 observability from the start.
        suspicious=[]
        req_headers={}
        def on_req(params):
            url=params.get('request',{}).get('url','')
            if SUSPICIOUS.search(url):
                req_headers[params.get('requestId')]=params
                suspicious.append({'event':'requestWillBeSent','ts':time.time(),'requestId':params.get('requestId'),'url':url,'method':params.get('request',{}).get('method'),'headers':params.get('request',{}).get('headers',{})})
        def on_resp(params):
            resp=params.get('response',{}); url=resp.get('url','')
            headers=resp.get('headers',{}) or {}
            hd=' '.join(f'{k}:{v}' for k,v in headers.items())
            if SUSPICIOUS.search(url) or re.search(r'content-disposition.*(attachment|docx|filename)', hd, re.I):
                suspicious.append({'event':'responseReceived','ts':time.time(),'requestId':params.get('requestId'),'url':url,'status':resp.get('status'),'mimeType':resp.get('mimeType'),'headers':headers})
        page_cdp.on('Network.requestWillBeSent', on_req)
        page_cdp.on('Network.responseReceived', on_resp)
        await page_cdp.send('Network.enable')

        cands=await find_candidates(page)
        log['initial_candidates']=[{k:v for k,v in c.items() if k!='handle'} for c in cands[:30]]
        (LOG_PATH).write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')

        # Tactic 1: raw Input.dispatchMouseEvent on export/docx candidates.
        t1={'name':'tactic1_browser_download_raw_mouse','started':time.strftime('%Y-%m-%dT%H:%M:%S')}
        log['tactics'].append(t1)
        for c in cands[:8]:
            await page.bring_to_front()
            t1.setdefault('tried',[]).append({k:v for k,v in c.items() if k!='handle'})
            await raw_click(page_cdp, c['bbox'], t1, f"t1 score={c['score']} {c['meta'].get('aria') or c['meta'].get('text')}")
            await page.wait_for_timeout(2500)
            await screenshot(page, f"01-t1-after-{len(t1['tried'])}.png")
            # If menu appeared, click Word/DOCX item.
            await click_format_items(page, page_cdp, t1)
            dl=await wait_for_download(download_state, timeout=12)
            t1.setdefault('download_waits',[]).append(dl)
            if dl:
                src=Path(dl.get('expected_path') or (RAW/dl.get('guid','')))
                if await copy_if_valid(src, log, 'tactic1'):
                    t1['worked']=True; break
            try: await page.keyboard.press('Escape')
            except Exception: pass
        if not t1.get('worked'):
            t1['worked']=False

        # Tactic 2: if suspicious response bodies/URLs exist, try fetching response body or refetching with cookies.
        log['network_suspicious_after_t1']=suspicious[-100:]
        if not looks_docx(FINAL).get('substantive'):
            t2={'name':'tactic2_network_response_refetch','started':time.strftime('%Y-%m-%dT%H:%M:%S')}
            log['tactics'].append(t2)
            # Click best candidate once more while network is enabled.
            cands=await find_candidates(page)
            if cands:
                await raw_click(page_cdp, cands[0]['bbox'], t2, 't2-network-best')
                await page.wait_for_timeout(3000)
                await click_format_items(page,page_cdp,t2)
                await page.wait_for_timeout(5000)
            t2['suspicious_count']=len(suspicious)
            t2['suspicious_tail']=suspicious[-80:]
            # Try Network.getResponseBody for docx-ish responses.
            for rec in reversed(suspicious[-120:]):
                hdr=' '.join(f"{k}:{v}" for k,v in (rec.get('headers') or {}).items())
                url=rec.get('url','')
                if rec.get('event')=='responseReceived' and re.search(r'(docx|attachment|filename|download|export)', url+' '+hdr, re.I):
                    rid=rec.get('requestId')
                    try:
                        body=await page_cdp.send('Network.getResponseBody', {'requestId':rid})
                        data=base64.b64decode(body['body']) if body.get('base64Encoded') else body.get('body','').encode()
                        out=RAW/(f'tactic2-body-{len(t2.get("bodies",[]))}.docx')
                        out.write_bytes(data)
                        v=looks_docx(out)
                        t2.setdefault('bodies',[]).append({'requestId':rid,'url':url,'path':str(out),'verify':v})
                        if v.get('substantive') and not v.get('byte_identical_to_manual'):
                            shutil.copy2(out, FINAL); shutil.copy2(out, CANON); t2['worked']=True; break
                    except Exception as e:
                        t2.setdefault('body_errors',[]).append({'requestId':rid,'url':url,'error':repr(e)[:400]})
            t2.setdefault('worked', False)

        # Tactic 3: Runtime/call .click inside iframe context.
        if not looks_docx(FINAL).get('substantive'):
            t3={'name':'tactic3_runtime_click_in_frame','started':time.strftime('%Y-%m-%dT%H:%M:%S')}
            log['tactics'].append(t3)
            for i,fr in enumerate(page.frames):
                try:
                    count=await fr.locator('button[aria-label="导出"], a:has-text("下载 DOCX"), a:has-text("DOCX")').count()
                except Exception as e:
                    t3.setdefault('frame_errors',[]).append({'frame':i,'error':repr(e)[:200]}); continue
                for j in range(min(count,3)):
                    try:
                        h=await fr.locator('button[aria-label="导出"], a:has-text("下载 DOCX"), a:has-text("DOCX")').nth(j).element_handle()
                        meta=await h.evaluate('el=>({text:el.innerText||el.textContent||"", aria:el.getAttribute("aria-label")||"", href:el.getAttribute("href")||""})')
                        t3.setdefault('tried',[]).append({'frame':i,'nth':j,'meta':meta})
                        await h.evaluate('el => el.click()')
                        await page.wait_for_timeout(2500)
                        await click_format_items(page,page_cdp,t3)
                        dl=await wait_for_download(download_state, timeout=10)
                        if dl:
                            if await copy_if_valid(Path(dl.get('expected_path') or (RAW/dl.get('guid',''))), log, 'tactic3'):
                                t3['worked']=True; break
                    except Exception as e:
                        t3.setdefault('errors',[]).append({'frame':i,'nth':j,'error':repr(e)[:300]})
                if t3.get('worked'): break
            t3.setdefault('worked', False)

        # Tactic 4: focus button and keyboard enter.
        if not looks_docx(FINAL).get('substantive'):
            t4={'name':'tactic4_focus_enter','started':time.strftime('%Y-%m-%dT%H:%M:%S')}
            log['tactics'].append(t4)
            cands=await find_candidates(page)
            for c in cands[:6]:
                try:
                    await c['handle'].evaluate('el => el.focus()')
                    t4.setdefault('tried',[]).append({k:v for k,v in c.items() if k!='handle'})
                    await page_cdp.send('Input.dispatchKeyEvent', {'type':'rawKeyDown','windowsVirtualKeyCode':13,'nativeVirtualKeyCode':13,'key':'Enter','code':'Enter'})
                    await page_cdp.send('Input.dispatchKeyEvent', {'type':'keyUp','windowsVirtualKeyCode':13,'nativeVirtualKeyCode':13,'key':'Enter','code':'Enter'})
                    await page.wait_for_timeout(2500)
                    await click_format_items(page,page_cdp,t4)
                    dl=await wait_for_download(download_state, timeout=10)
                    if dl and await copy_if_valid(Path(dl.get('expected_path') or (RAW/dl.get('guid',''))), log, 'tactic4'):
                        t4['worked']=True; break
                except Exception as e:
                    t4.setdefault('errors',[]).append({'error':repr(e)[:300]})
            t4.setdefault('worked', False)

        # Tactic 5: Fetch interception at response stage for attachment/docx.
        if not looks_docx(FINAL).get('substantive'):
            t5={'name':'tactic5_fetch_response_body','started':time.strftime('%Y-%m-%dT%H:%M:%S')}
            log['tactics'].append(t5)
            paused=[]
            async def handle_fetch(params):
                rid=params.get('requestId')
                url=params.get('request',{}).get('url','')
                hdrs=params.get('responseHeaders') or []
                hd=' '.join(f"{h.get('name')}:{h.get('value')}" for h in hdrs)
                matched=bool(re.search(r'(content-disposition.*(attachment|docx|filename)|\.docx|application/vnd\.openxmlformats)', hd+' '+url, re.I))
                rec={'url':url,'requestId':rid,'matched':matched,'headers':hdrs[:20]}
                paused.append(rec)
                try:
                    if matched:
                        body=await page_cdp.send('Fetch.getResponseBody', {'requestId':rid})
                        data=base64.b64decode(body['body']) if body.get('base64Encoded') else body.get('body','').encode()
                        out=RAW/f'tactic5-fetch-{len(paused)}.docx'
                        out.write_bytes(data)
                        rec['saved']=str(out); rec['verify']=looks_docx(out)
                        if rec['verify'].get('substantive') and not rec['verify'].get('byte_identical_to_manual'):
                            shutil.copy2(out, FINAL); shutil.copy2(out, CANON); t5['worked']=True
                    await page_cdp.send('Fetch.continueRequest', {'requestId':rid})
                except Exception as e:
                    rec['error']=repr(e)[:300]
                    try: await page_cdp.send('Fetch.continueRequest', {'requestId':rid})
                    except Exception: pass
            page_cdp.on('Fetch.requestPaused', lambda params: asyncio.create_task(handle_fetch(params)))
            try:
                await page_cdp.send('Fetch.enable', {'patterns':[{'urlPattern':'*','requestStage':'Response'}]})
                cands=await find_candidates(page)
                if cands:
                    await raw_click(page_cdp, cands[0]['bbox'], t5, 't5-fetch-best')
                    await page.wait_for_timeout(2500)
                    await click_format_items(page,page_cdp,t5)
                    await page.wait_for_timeout(10000)
                await page_cdp.send('Fetch.disable')
            except Exception as e:
                t5['error']=repr(e)
                try: await page_cdp.send('Fetch.disable')
                except Exception: pass
            t5['paused_tail']=paused[-80:]
            t5.setdefault('worked', False)

        log['final_screenshot']=await screenshot(page,'99-final.png', full=False)
        log['final_verify']=looks_docx(FINAL)
        log['raw_files']=[{'name':p.name,'size':p.stat().st_size,'sha256':sha256(p)} for p in RAW.glob('*') if p.is_file()]
        log['network_suspicious_final']=suspicious[-150:]
        LOG_PATH.write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
        print(json.dumps(log['final_verify'], ensure_ascii=False, indent=2))
        print('LOG', LOG_PATH)

if __name__ == '__main__':
    asyncio.run(main())
