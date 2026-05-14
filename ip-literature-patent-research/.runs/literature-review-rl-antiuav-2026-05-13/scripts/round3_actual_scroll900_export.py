import asyncio, json, time, shutil, hashlib, re, base64
from pathlib import Path
from playwright.async_api import async_playwright
RUN=Path('ip-literature-patent-research/.runs/literature-review-rl-antiuav-2026-05-13').resolve(); RAW=RUN/'round3-raw-downloads'; SHOTS=RUN/'round3-screenshots'; SCRIPTS=RUN/'scripts'; RAW.mkdir(exist_ok=True); SHOTS.mkdir(exist_ok=True)
URL='https://chatgpt.com/c/6a04a213-5648-83e8-b9d0-6134aef56831'; FINAL=RUN/'round3-chatgpt-dr-export-strong-form.docx'; CANON=RUN/'强化学习在反无人机系统中的应用-文献综述.docx'; MANUAL=Path.home()/'Downloads'/'强化学习在反无人机系统中的应用.docx'
def sha(p):
 if not p.exists(): return None
 h=hashlib.sha256(); h.update(p.read_bytes()); return h.hexdigest()
def verify(p):
 r={'exists':p.exists()}
 if p.exists():
  r.update(size=p.stat().st_size, sha256=sha(p), manual_sha256=sha(MANUAL)); r['byte_identical_to_manual']=r['sha256']==r['manual_sha256']
  try:
   from docx import Document
   d=Document(str(p)); r['paragraphs']=len(d.paragraphs); r['chars']=sum(len(x.text) for x in d.paragraphs); r['valid_docx']=True; r['substantive']=r['paragraphs']>=50 and r['chars']>=5000
  except Exception as e: r['valid_docx']=False; r['error']=repr(e)
 return r
async def raw(cdp,b):
 cx=b['x']+b['width']/2; cy=b['y']+b['height']/2
 for ev in [{'type':'mouseMoved','x':cx,'y':cy},{'type':'mousePressed','x':cx,'y':cy,'button':'left','buttons':1,'clickCount':1},{'type':'mouseReleased','x':cx,'y':cy,'button':'left','buttons':0,'clickCount':1}]: await cdp.send('Input.dispatchMouseEvent',ev)
 return {'x':cx,'y':cy}
async def set_main_scroll(page, top):
 await page.evaluate('''top=>{const els=Array.from(document.querySelectorAll('*')).filter(el=>el.scrollHeight>el.clientHeight+50 && getComputedStyle(el).overflowY==='auto'); const main=els.find(el=>el.getBoundingClientRect().x>200 && el.clientHeight>900); if(main) main.scrollTop=top;}''', top)
async def main():
 log={'started':time.strftime('%FT%T'),'manual_sha256':sha(MANUAL)}
 async with async_playwright() as p:
  b=await p.chromium.connect_over_cdp('http://127.0.0.1:9223'); ctx=b.contexts[0]
  page=await ctx.new_page(); await page.goto(URL, wait_until='domcontentloaded', timeout=60000); await page.set_viewport_size({'width':1500,'height':1000}); await page.wait_for_timeout(15000); await page.bring_to_front()
  cdp=await ctx.new_cdp_session(page); bcdp=await b.new_browser_cdp_session(); downloads={}; events=[]; suspicious=[]
  bcdp.on('Browser.downloadWillBegin', lambda e:(downloads.setdefault(e.get('guid'),{}).update(e), events.append(dict(e,event='will',ts=time.time()))))
  bcdp.on('Browser.downloadProgress', lambda e:(downloads.setdefault(e.get('guid'),{}).update(e), events.append(dict(e,event='progress',ts=time.time()))))
  await bcdp.send('Browser.setDownloadBehavior', {'behavior':'allowAndName','downloadPath':str(RAW),'eventsEnabled':True})
  await cdp.send('Network.enable')
  cdp.on('Network.responseReceived', lambda e: suspicious.append({'requestId':e.get('requestId'),'url':e.get('response',{}).get('url'),'status':e.get('response',{}).get('status'),'headers':e.get('response',{}).get('headers',{})}) if re.search(r'docx|download|export|attachment|file|conversation', e.get('response',{}).get('url','')+' '+str(e.get('response',{}).get('headers',{})), re.I) else None)
  await set_main_scroll(page, 900); await page.wait_for_timeout(1000); await page.screenshot(path=str(SHOTS/'actual-scroll900-before.png'))
  # locate full report export by context
  target=None
  for i,fr in enumerate(page.frames):
   hs=await fr.locator('button[aria-label="导出"]').element_handles()
   for j,h in enumerate(hs):
    ctxs=await h.evaluate('el=>{let p=el,ctx="";for(let i=0;p&&i<5;i++,p=p.parentElement)ctx+=(p.innerText||p.textContent||"").slice(0,800);return ctx}')
    box=await h.bounding_box()
    if box and 0<=box['y']<=1000 and '引言与背景' in ctxs:
     target={'frame':i,'nth':j,'ctx':ctxs[:1000],'bbox':box,'handle':h}
  log['target']={k:v for k,v in (target or {}).items() if k!='handle'}
  if target:
   log['click']=await raw(cdp,target['bbox']); await page.wait_for_timeout(2500); await page.screenshot(path=str(SHOTS/'actual-scroll900-after-export.png'))
   # don't click old attachment links unless visible menu item is near target and has Word/DOCX
   menu=[]
   for i,fr in enumerate(page.frames):
    hs=await fr.locator('[role="menuitem"], button, a').element_handles()
    for j,h in enumerate(hs):
     try:
      txt=await h.evaluate('el=>((el.innerText||el.textContent||"")+" "+(el.getAttribute("aria-label")||"")+" "+(el.getAttribute("href")||"")).trim()')
      box=await h.bounding_box()
      if box and 0<=box['y']<=1000 and re.search(r'DOCX|Word|下载 DOCX|导出|Download', txt, re.I): menu.append({'frame':i,'nth':j,'text':txt[:300],'bbox':box,'handle':h})
     except Exception: pass
   log['menu_items']=[{k:v for k,v in m.items() if k!='handle'} for m in menu]
   # click docx/word menu item if it is not an old attachment far below and appeared near button
   for m in menu:
    if re.search(r'DOCX|Word', m['text'], re.I) and abs(m['bbox']['y']-target['bbox']['y'])<250:
     log['menu_click']=await raw(cdp,m['bbox']); await page.wait_for_timeout(2500); break
  for _ in range(60):
   comp=[(g,d) for g,d in downloads.items() if d.get('state')=='completed']; files=[f for f in RAW.iterdir() if f.is_file() and f.stat().st_size>1000 and not f.name.endswith('.crdownload')]
   if comp or files:
    src=RAW/comp[-1][0] if comp else max(files,key=lambda f:f.stat().st_mtime)
    if not src.exists() and files: src=max(files,key=lambda f:f.stat().st_mtime)
    shutil.copy2(src,FINAL); log['src']=str(src); log['verify']=verify(FINAL)
    if log['verify'].get('substantive') and not log['verify'].get('byte_identical_to_manual'): shutil.copy2(FINAL,CANON); log['canonical_overwritten']=True
    break
   await asyncio.sleep(.5)
  log['events']=events; log['suspicious']=suspicious[-100:]; log['final_verify']=verify(FINAL)
  (SCRIPTS/'round3-actual-scroll900-export-log.json').write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
  print(json.dumps(log,ensure_ascii=False,indent=2)[:8000])
asyncio.run(main())
