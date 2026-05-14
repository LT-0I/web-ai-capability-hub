import asyncio, json, time, shutil, hashlib, zipfile
from pathlib import Path
from playwright.async_api import async_playwright
RUN=Path('ip-literature-patent-research/.runs/literature-review-rl-antiuav-2026-05-13').resolve(); RAW=RUN/'round3-raw-downloads'; SHOTS=RUN/'round3-screenshots'; SCRIPTS=RUN/'scripts'
RAW.mkdir(exist_ok=True); SHOTS.mkdir(exist_ok=True)
URL='https://chatgpt.com/c/6a04a213-5648-83e8-b9d0-6134aef56831'
FINAL=RUN/'round3-chatgpt-dr-export-strong-form.docx'; CANON=RUN/'强化学习在反无人机系统中的应用-文献综述.docx'; MANUAL=Path.home()/'Downloads'/'强化学习在反无人机系统中的应用.docx'
LOG=SCRIPTS/'round3-focused-raw-click-log.json'

def sha(p):
 if not p.exists(): return None
 h=hashlib.sha256(); h.update(p.read_bytes()); return h.hexdigest()
def verify(p):
 r={'exists':p.exists(),'path':str(p)}
 if not p.exists(): return r
 r['size']=p.stat().st_size; r['sha256']=sha(p); r['manual_sha256']=sha(MANUAL); r['byte_identical_to_manual']=(r['sha256']==r['manual_sha256'])
 try:
  from docx import Document
  d=Document(str(p)); r['paragraphs']=len(d.paragraphs); r['chars']=sum(len(x.text) for x in d.paragraphs); r['valid_docx']=True
 except Exception as e: r['valid_docx']=False; r['error']=repr(e)
 r['substantive']=r.get('valid_docx') and r.get('paragraphs',0)>=50 and r.get('chars',0)>=5000
 return r
async def candidates(page):
 out=[]
 for i,fr in enumerate(page.frames):
  for sel in ['button[aria-label="导出"]','a:has-text("下载 DOCX")','a:has-text("DOCX")','button,a']:
   try: hs=await fr.locator(sel).element_handles()
   except Exception: hs=[]
   for j,h in enumerate(hs):
    try:
     meta=await h.evaluate('''el=>{const clean=s=>(s||'').replace(/\s+/g,' ').trim(); const r=el.getBoundingClientRect(); let p=el,ctx=''; for(let i=0;p&&i<4;i++,p=p.parentElement) ctx+=' || '+clean(p.innerText||p.textContent||''); return {tag:el.tagName.toLowerCase(), text:clean(el.innerText||el.textContent||''), aria:el.getAttribute('aria-label')||'', href:el.getAttribute('href')||'', box:{x:r.x,y:r.y,w:r.width,h:r.height}, ctx:ctx.slice(0,500)}}''')
     blob=(meta['text']+' '+meta['aria']+' '+meta['href']+' '+meta['ctx']).lower()
     score=0
     if meta['aria']=='导出': score+=100
     if '下载 docx' in blob: score+=80
     if 'docx' in blob: score+=50
     if '导出' in meta['aria'] or 'export' in blob: score+=20
     if 'pptx' in blob and 'docx' not in blob: score-=30
     if meta['box']['w'] and meta['box']['h'] and score>0:
      out.append({'score':score,'frame':i,'selector':sel,'nth':j,'meta':meta,'handle':h})
    except Exception: pass
 out.sort(key=lambda x:x['score'], reverse=True)
 return out
async def raw_click(cdp,bbox):
 cx=bbox['x']+bbox['width']/2; cy=bbox['y']+bbox['height']/2
 await cdp.send('Input.dispatchMouseEvent', {'type':'mouseMoved','x':cx,'y':cy})
 await cdp.send('Input.dispatchMouseEvent', {'type':'mousePressed','x':cx,'y':cy,'button':'left','buttons':1,'clickCount':1})
 await cdp.send('Input.dispatchMouseEvent', {'type':'mouseReleased','x':cx,'y':cy,'button':'left','buttons':0,'clickCount':1})
 return {'x':cx,'y':cy}
async def main():
 log={'started':time.strftime('%FT%T'),'manual_sha256':sha(MANUAL),'tries':[]}
 async with async_playwright() as p:
  b=await p.chromium.connect_over_cdp('http://127.0.0.1:9223'); ctx=b.contexts[0]
  page=None
  for pg in ctx.pages:
   if pg.url.split('?')[0]==URL: page=pg; break
  if not page: page=await ctx.new_page(); await page.goto(URL, wait_until='domcontentloaded', timeout=60000)
  await page.bring_to_front(); await page.set_viewport_size({'width':1500,'height':1000}); await page.wait_for_timeout(3000)
  cdp=await ctx.new_cdp_session(page); bcdp=await b.new_browser_cdp_session()
  downloads={}
  bcdp.on('Browser.downloadWillBegin', lambda e: (downloads.setdefault(e.get('guid'),{}).update(e), log.setdefault('download_events',[]).append(dict(e,event='willBegin',ts=time.time()))))
  bcdp.on('Browser.downloadProgress', lambda e: (downloads.setdefault(e.get('guid'),{}).update(e), log.setdefault('download_events',[]).append(dict(e,event='progress',ts=time.time()))))
  await bcdp.send('Browser.setDownloadBehavior', {'behavior':'allowAndName','downloadPath':str(RAW),'eventsEnabled':True})
  await page.screenshot(path=str(SHOTS/'focused-00.png'), full_page=True)
  cs=await candidates(page); log['initial_candidates']=[{k:v for k,v in c.items() if k!='handle'} for c in cs[:50]]; LOG.write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
  # try export buttons first, latest/lower buttons first by y after score
  cs.sort(key=lambda c:(c['score'], c['meta']['box']['y']), reverse=True)
  for idx,c in enumerate(cs[:20]):
   if c['score'] < 50: continue
   rec={k:v for k,v in c.items() if k!='handle'}; log['tries'].append(rec)
   try:
    await c['handle'].scroll_into_view_if_needed(timeout=5000); await page.wait_for_timeout(600)
    bbox=await c['handle'].bounding_box(); rec['bbox_after_scroll']=bbox
    await page.screenshot(path=str(SHOTS/f'focused-before-{idx}.png'), full_page=False)
    rec['click_point']=await raw_click(cdp,bbox)
    await page.wait_for_timeout(1500)
    # if menu/picker/link appears, click visible DOCX/Word links only
    cs2=await candidates(page); rec['post_candidates']=[{k:v for k,v in x.items() if k!='handle'} for x in cs2[:15]]
    for x in cs2[:10]:
     blob=(x['meta']['text']+' '+x['meta']['aria']+' '+x['meta']['ctx']).lower()
     if ('docx' in blob or 'word' in blob) and x['score']>=50:
      await x['handle'].scroll_into_view_if_needed(timeout=3000); await page.wait_for_timeout(300)
      bb=await x['handle'].bounding_box(); rec.setdefault('post_clicks',[]).append({'candidate':{k:v for k,v in x.items() if k!='handle'}, 'bbox':bb})
      await raw_click(cdp,bb); await page.wait_for_timeout(1500); break
    # wait max 20s
    for _ in range(40):
     completed=[(g,d) for g,d in downloads.items() if d.get('state')=='completed']
     files=[f for f in RAW.iterdir() if f.is_file() and f.stat().st_size>1000 and not f.name.endswith('.crdownload')]
     if completed or files:
      src=None
      if completed: src=RAW/completed[-1][0]; rec['completed_event']=completed[-1][1]
      if not src or not src.exists(): src=max(files,key=lambda f:f.stat().st_mtime)
      rec['src']=str(src); shutil.copy2(src, FINAL); rec['verify']=verify(FINAL)
      if rec['verify'].get('substantive') and not rec['verify'].get('byte_identical_to_manual'):
       shutil.copy2(FINAL,CANON); rec['canonical_overwritten']=True; log['worked']=True; break
     await asyncio.sleep(0.5)
    LOG.write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
    if log.get('worked'): break
    await page.keyboard.press('Escape'); await page.wait_for_timeout(500)
   except Exception as e:
    rec['error']=repr(e); LOG.write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
  log['final_verify']=verify(FINAL); log['raw_files']=[{'name':f.name,'size':f.stat().st_size,'sha256':sha(f)} for f in RAW.iterdir() if f.is_file()]
  await page.screenshot(path=str(SHOTS/'focused-99.png'), full_page=False)
  LOG.write_text(json.dumps(log,ensure_ascii=False,indent=2),encoding='utf-8')
  print(json.dumps(log.get('final_verify'),ensure_ascii=False,indent=2))
asyncio.run(main())
